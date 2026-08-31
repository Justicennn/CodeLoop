"""Stage 10A local PDF/DOCX source adapter tests."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

from codeloop.execution.tools import MAX_DOCUMENT_CHARS, ToolRegistry
from codeloop.execution.workspace import Workspace


def _dispatch(
    registry: ToolRegistry,
    path: str,
    **arguments: object,
) -> dict[str, object]:
    return registry.dispatch(
        "read_document",
        json.dumps({"path": path, **arguments}),
    )


def _write_text_pdf(path: Path, pages: list[str | None]) -> None:
    """Build a small deterministic text PDF without a generator dependency."""
    writer = PdfWriter()
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_reference = writer._add_object(font)
    for text in pages:
        page = writer.add_blank_page(width=300, height=300)
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): font_reference}
                )
            }
        )
        if text is not None:
            stream = DecodedStreamObject()
            stream.set_data(
                f"BT /F1 12 Tf 20 200 Td ({text}) Tj ET".encode("ascii")
            )
            page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)


def test_pdf_text_order_normalization_and_cursor_contract(tmp_path: Path) -> None:
    path = tmp_path / "requirements.pdf"
    _write_text_pdf(path, ["Alpha", "Beta"])
    registry = ToolRegistry(Workspace(tmp_path))
    before = path.read_bytes()

    first = _dispatch(registry, path.name, cursor=0, max_chars=7)
    assert first["ok"] is True
    assert first["data"] == {
        "path": "requirements.pdf",
        "document_type": "pdf",
        "text": "Alpha\n\n",
        "position": {
            "start_cursor": 0,
            "end_cursor": 7,
            "total_chars": 11,
            "first_unit": {"kind": "page", "index": 1},
            "last_unit": {"kind": "page", "index": 1},
        },
        "truncated": True,
        "next_cursor": 7,
    }

    second = _dispatch(registry, path.name, cursor=7, max_chars=10)
    assert second["ok"] is True
    assert second["data"]["text"] == "Beta"
    assert second["data"]["position"]["first_unit"] == {
        "kind": "page",
        "index": 2,
    }
    assert second["data"]["truncated"] is False
    assert second["data"]["next_cursor"] is None

    eof = _dispatch(registry, path.name, cursor=11)
    assert eof["ok"] is True
    assert eof["data"]["text"] == ""
    assert eof["data"]["position"] == {
        "start_cursor": 11,
        "end_cursor": 11,
        "total_chars": 11,
        "first_unit": None,
        "last_unit": None,
    }
    assert eof["data"]["truncated"] is False
    assert eof["data"]["next_cursor"] is None

    beyond = _dispatch(registry, path.name, cursor=12)
    assert beyond["ok"] is False
    assert beyond["error_code"] == "invalid_arguments"
    assert beyond["data"] == {"cursor": 12, "total_chars": 11}
    assert path.read_bytes() == before


def test_pdf_empty_units_determinism_and_no_text_error(tmp_path: Path) -> None:
    mixed = tmp_path / "mixed.pdf"
    _write_text_pdf(mixed, [None, "Visible"])
    registry = ToolRegistry(Workspace(tmp_path))

    observed = _dispatch(registry, mixed.name)
    repeated = _dispatch(registry, mixed.name)
    assert observed == repeated
    assert observed["data"]["text"] == "\n\nVisible"
    assert observed["data"]["position"]["first_unit"] == {
        "kind": "page",
        "index": 1,
    }

    blank = tmp_path / "blank.pdf"
    _write_text_pdf(blank, [None])
    unavailable = _dispatch(registry, blank.name)
    assert unavailable["ok"] is False
    assert unavailable["error_code"] == "document_text_unavailable"


def test_malformed_pdf_is_a_structured_error(tmp_path: Path) -> None:
    (tmp_path / "broken.pdf").write_bytes(b"not a pdf")
    result = _dispatch(ToolRegistry(Workspace(tmp_path)), "broken.pdf")
    assert result["ok"] is False
    assert result["error_code"] == "malformed_document"


def test_docx_paragraph_table_unicode_and_block_order(tmp_path: Path) -> None:
    path = tmp_path / "requirements.docx"
    document = Document()
    document.add_paragraph("  功能要求\r第一项  ")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "名称"
    table.cell(0, 1).text = "数量"
    table.cell(1, 0).text = "商品"
    table.cell(1, 1).text = "2"
    document.add_paragraph("验收完成")
    document.save(path)

    before = path.read_bytes()
    result = _dispatch(ToolRegistry(Workspace(tmp_path)), path.name)
    after = path.read_bytes()

    assert result["ok"] is True
    assert result["data"]["document_type"] == "docx"
    assert result["data"]["text"] == (
        "功能要求\n第一项\n"
        "名称\t数量\n商品\t2\n"
        "验收完成"
    )
    assert result["data"]["position"]["first_unit"] == {
        "kind": "paragraph",
        "index": 1,
    }
    assert result["data"]["position"]["last_unit"] == {
        "kind": "paragraph",
        "index": 3,
    }
    assert before == after


def test_docx_cursor_can_start_inside_table_block(tmp_path: Path) -> None:
    path = tmp_path / "table.docx"
    document = Document()
    document.add_paragraph("P")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    document.save(path)
    registry = ToolRegistry(Workspace(tmp_path))

    result = _dispatch(registry, path.name, cursor=2, max_chars=2)
    assert result["ok"] is True
    assert result["data"]["text"] == "A\t"
    assert result["data"]["position"]["first_unit"] == {
        "kind": "table",
        "index": 2,
    }


def test_malformed_docx_and_unsupported_extension(tmp_path: Path) -> None:
    (tmp_path / "broken.docx").write_bytes(b"not a docx")
    (tmp_path / "requirements.txt").write_text("Use read_file", encoding="utf-8")
    registry = ToolRegistry(Workspace(tmp_path))

    malformed = _dispatch(registry, "broken.docx")
    assert malformed["ok"] is False
    assert malformed["error_code"] == "malformed_document"

    unsupported = _dispatch(registry, "requirements.txt")
    assert unsupported["ok"] is False
    assert unsupported["error_code"] == "unsupported_document_type"


def test_document_arguments_and_workspace_containment(tmp_path: Path) -> None:
    path = tmp_path / "requirements.pdf"
    _write_text_pdf(path, ["inside"])
    outside = tmp_path.parent / f"{tmp_path.name}-outside.pdf"
    _write_text_pdf(outside, ["outside"])
    registry = ToolRegistry(Workspace(tmp_path))

    for value in ("../outside.pdf", str(outside)):
        result = _dispatch(registry, value)
        assert result["ok"] is False
        assert result["error_code"] == "invalid_path"

    escape = tmp_path / "escape.pdf"
    try:
        escape.symlink_to(outside)
    except (NotImplementedError, OSError):
        pass
    else:
        result = _dispatch(registry, escape.name)
        assert result["ok"] is False
        assert result["error_code"] == "invalid_path"

    for arguments in (
        {"cursor": -1},
        {"max_chars": 0},
        {"max_chars": MAX_DOCUMENT_CHARS + 1},
    ):
        result = _dispatch(registry, path.name, **arguments)
        assert result["ok"] is False
        assert result["error_code"] == "invalid_arguments"
